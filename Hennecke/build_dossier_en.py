#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Build the English translation of the HENNECKE-OMS PE dossier, embedding the
English corporate-structure board chart in section 2.4."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "/home/user/Storage/Hennecke/Dossier_HenneckeOMS_PE_EN.docx"
CHART = "/home/user/Storage/Hennecke/corporate-structure-board.png"

doc = Document()

# ---------- helpers ----------
def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)
def p(t=""):
    para = doc.add_paragraph(t)
    return para
def li(t): doc.add_paragraph(t, style="List Bullet")

def table(header, rows):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    hc = t.rows[0].cells
    for i, txt in enumerate(header):
        hc[i].text = str(txt)
        for par in hc[i].paragraphs:
            for run in par.runs:
                run.font.bold = True
    for row in rows:
        rc = t.add_row().cells
        for i, txt in enumerate(row):
            rc[i].text = str(txt)
    return t

# ---------- Title block ----------
title = doc.add_heading("HENNECKE-OMS S.p.A.", level=0)
sub = p("Analysis dossier – Corporate, industrial and financial profile")
sub.runs[0].italic = True
p("Period analysed: Financial years 2023 – 2024 (filed financial statements)")
p("Reference sector: Machines and plants for polyurethane (ATECO 28.96.00 – NACE 2829)")
note = p("Confidential document. Information drawn from: Coface Business Information – Full Report "
         "(produced on 04/06/2026); Historical company extract (visura), Chamber of Commerce of "
         "Milano-Monza-Brianza-Lodi (extracted on 04/06/2026); public industry sources (official "
         "Hennecke Group website, trade press). The consolidated financial statements of the Hennecke "
         "Group (Germany) were not part of this analysis.")
note.runs[0].font.size = Pt(9)

p("Note on figures: amounts are in € thousand unless otherwise stated. Numbers follow the European "
  "format of the source filings (\".\" = thousands separator, \",\" = decimal separator).").runs[0].italic = True

# ---------- Index ----------
h1("Index")
for line in [
    "1.  Executive Summary",
    "2.  Corporate Profile, Governance & Structure",
    "3.  Products, Technology & Commercial Positioning",
    "4.  Operations & Industrial Footprint",
    "5.  Market Analysis – Polyurethane Machinery",
    "6.  Competitor Mapping",
    "7.  Financial Statement Analysis 2023-2024",
    "8.  Earnings Dynamics & Variance Analysis",
    "9.  Considerations for Private Equity",
    "10. Areas for Further Due Diligence",
    "11. SWOT Analysis",
    "12. Sources & Disclaimer",
    "13. Appendix – Financial Schedules (extended format)",
]:
    p(line)

# ---------- 1. Executive Summary ----------
h1("1. Executive Summary")
p("HENNECKE-OMS S.p.A. is the Italian production hub of the German Hennecke group, one of the world's "
  "leading manufacturers of machines and plants for polyurethane processing. Headquartered with its plant "
  "in Verano Brianza (MB), the company has its roots in the historic OMS (Officine Meccaniche) of the "
  "Brianza capital-goods engineering district, progressively integrated into the Hennecke group between "
  "2018 and 2020 (sole shareholder Hennecke GmbH since December 2020). It designs, manufactures, commissions "
  "and markets complete lines for flexible and rigid polyurethane foams: slabstock, continuous and "
  "discontinuous insulating sandwich panels, refrigeration, automotive, footwear, elastomers and composites. "
  "In 2024 it generated revenue of about €44.2 million with an average workforce of ~166 employees.")

h2("1.1 Key figures 2023-2024 (€ thousand unless otherwise stated)")
table(["Indicator", "2023", "2024", "Δ YoY"], [
    ["Revenue (sales & services)", "50.389", "44.243", "-12,2%"],
    ["Value of production", "51.998", "40.652", "-21,8%"],
    ["EBITDA (Coface)", "3.624", "1.542", "-57,5%"],
    ["EBITDA margin", "7,2%", "3,5%", "-3,7 p.p."],
    ["EBIT", "1.015", "-1.007", "n.m."],
    ["Net profit (loss) for the year", "289", "-241", "n.m."],
    ["Shareholders' equity", "6.137", "6.043", "-1,5%"],
    ["Intangible assets", "13.348", "12.016", "-10,0%"],
    ["Cash and cash equivalents", "8.371", "10.365", "+23,8%"],
    ["Bank debt (medium/long-term)", "0", "6.907", "n.m."],
    ["Customer advances/prepayments", "14.319", "10.948", "-23,5%"],
    ["Cash flow", "2.898", "2.308", "-20,4%"],
    ["Employees (average)", "170", "166", "-4 (159 in 2025)"],
])

h2("1.2 Three key takeaways")
p("Project-based, cyclical business tied to capital goods.")
p("2024 shows revenue -12% and, above all, value of production -22%: the divergence stems from the "
  "long project-cycle nature of the business and the dynamics of work in progress (WIP change of +€19.3M "
  "in 2023 versus -€0.5M in 2024). EBITDA compresses to €1.5M (3.5%), with a net loss of €0.24M. "
  "Profitability \"beneath the surface\" reflects order timing more than structural deterioration, but "
  "remains thin.")
p("Captive asset of the Hennecke group.")
p("The company is 100%-controlled by Hennecke GmbH (Sankt Augustin, Germany), is subject to direction and "
  "coordination and operates within the group's \"Financial Playbook\" (signature thresholds, treasury, "
  "intercompany orders). The shares are pledged. Any Private Equity transaction would take the form of a "
  "carve-out of the Italian entity or would involve the entire Hennecke Group (in turn attributable to a "
  "German industrial holding).")
p("Distinctive financial structure: low debt, high liquidity, self-financed working capital.")
p("Equity is thin (€6.0M, 15.5% of total assets); the balance sheet is weighed down by intangible assets "
  "of €12.0M (goodwill/intangibles from business combinations). Operating working capital is largely "
  "financed by customer advances (€10.9M) and liquidity is high (€10.4M) against a contained medium/long-term "
  "bank debt (€6.9M). The Coface credit score is 5/10 (medium-high insolvency risk), with a recommended "
  "maximum exposure of €500k.")

h2("1.3 Investment thesis summary")
p("Rationale: a polyurethane plant-engineering platform with a historic brand (OMS), a global installed base "
  "and full integration into a world-leading group; distinctive competencies in continuous sandwich-panel "
  "lines and flexible foams. Value-creation levers: (i) normalisation of the order book and recovery of "
  "margins towards 6-8% (fixed-cost efficiency, leverage on service and spare parts); (ii) development of the "
  "high-margin, recurring aftermarket (spare parts, retrofit, service); (iii) exposure to structural drivers "
  "(building energy efficiency, cold chain, e-mobility, blowing-agent transition); (iv) working-capital "
  "rationalisation. Main critical issues: thin and volatile margins, dependence on the customers' investment "
  "cycle, captive nature (limited autonomy), reduced equity and high intangible assets to be tested for "
  "impairment.")

# ---------- 2. Corporate Profile ----------
h1("2. Corporate Profile, Governance & Structure")
h2("2.1 Company details")
table(["Item", "Detail"], [
    ["Company name", "HENNECKE-OMS S.P.A."],
    ["Registered office", "Via Sabbionetta 4, 20843 Verano Brianza (MB)"],
    ["Tax code / VAT no.", "08322500151"],
    ["REA (company register no.)", "MB - 1218538"],
    ["LEI code", "815600D4EF31E2D24E18"],
    ["Legal form", "Joint-stock company (Società per azioni)"],
    ["Share capital (fully paid-in)", "€ 2.322.000,00 (450.000 shares)"],
    ["Incorporation date", "18/04/1986 (orig. \"Fiorio Services S.p.A.\")"],
    ["Duration", "until 31/12/2050"],
    ["ATECO / NACE", "28.96.00 / 28.96 – Machinery for plastics & rubber industry"],
    ["Corporate purpose", "Design, manufacture, commissioning and sale of machines and plants for the foaming of expanded polyurethane"],
    ["Control", "100% Hennecke GmbH (DE); subject to direction and coordination"],
    ["Auditor", "KPMG S.p.A."],
    ["Employees (average)", "159 (2025) · 166 (2024) · 170 (2023)"],
    ["Certifications", "UNI EN ISO 9001:2015 (since 2011) · UNI EN ISO 14001:2015 (since 2024), TÜV Italia"],
    ["Coface credit score", "5/10 (medium-high insolvency risk)"],
])

h2("2.2 Governance (Board in office, appointments 28/04/2026 and 04/05/2026)")
p("In April 2026 Andrea Enrico Carlo Mariani — the last member of the OMS founding family on the Board — "
  "left office: the renewal marks the completion of the governance handover to the control of the German "
  "group, with a board composed of a majority of managers representing Hennecke and the arrival of an "
  "Italian Managing Director (Mari) with operational powers.")
table(["Role", "Person", "Notes"], [
    ["Chairman of the Board", "Rolf Trippler", "Born in Leverkusen (DE), 1962; company representative"],
    ["Managing Director (CEO)", "Thomas Wildt", "Born in Frechen (DE), 1965"],
    ["Managing Director", "Alessandro Walter Mari", "Born in Milan, 1978; MD since 04/05/2026 (formerly proxy)"],
    ["Board of Statutory Auditors", "R. Lazzarone (Chair), M. Pietra, R. De Bernardinis", "Alternates: G. Pistillo, R. Patella"],
    ["Independent auditor", "KPMG S.p.A.", "Appointed until FY2028 financial statements"],
    ["Special attorneys (proxies)", "S. Sala, R. Diaferio, G.C. Fumagalli, P. Masciulli", "Banking/operational powers, joint signature"],
])

h2("2.3 Shareholding")
p("HENNECKE-OMS is a wholly-owned company. The capital (450,000 shares, €2,322,000) is 100%-held by "
  "Hennecke GmbH, based in Sankt Augustin (Germany), recorded as sole shareholder since 15/12/2020. The "
  "stake consists of pledged shares (a constraint to be verified during Due Diligence, likely as collateral "
  "for group financing lines). The company prepares a declaration of membership of the Hennecke group and is "
  "subject to a third party's direction and coordination (art. 2497-bis of the Italian Civil Code).")
table(["Shareholder", "Stake", "Nature"], [
    ["Hennecke GmbH (Sankt Augustin, DE)", "100%", "Sole shareholder / parent company (since 12/2020)"],
])

h2("2.4 Structure and membership of the Hennecke Group")
p("The company is the Italian division of the Hennecke Group, one of the world leaders in process "
  "technologies for polyurethane. The parent Hennecke GmbH (Germany) oversees R&D, the international sales "
  "network and the global installed base; HENNECKE-OMS contributes the historic OMS know-how on flexible-foam "
  "plants and sandwich-panel lines. The group operates with a centralised governance model (common signature "
  "thresholds and \"Financial Playbook\", purchase orders to group suppliers, coordinated treasury), which "
  "entails a significant component of intercompany relationships (commercial and financial) to be isolated "
  "during the analysis.")
p("The perimeter and the ownership arrangements upstream of Hennecke GmbH (German industrial holding) are "
  "inferred from public industry sources and not from official documentation analysed.")
# --- embed the English corporate-structure chart ---
doc.add_paragraph()
cap = doc.add_paragraph()
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap_run = cap.add_run("Hennecke GROUP — new corporate structure (following the acquisition by Brückner Group SE)")
cap_run.italic = True
cap_run.font.size = Pt(9)
pic_par = doc.add_paragraph()
pic_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
pic_par.add_run().add_picture(CHART, width=Inches(6.3))
src = doc.add_paragraph()
src.alignment = WD_ALIGN_PARAGRAPH.CENTER
src_run = src.add_run("Note: on 5 Jan 2026 the Hennecke Group was acquired by Brückner Group SE from a fund "
                      "advised by Capvis AG. Dates from press sources, to be confirmed against official documentation.")
src_run.font.size = Pt(8)
src_run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

h2("2.5 Industrial history")
li("1986 – Incorporation of the company (originally \"Fiorio Services S.p.A.\").")
li("1988 – Merger of Impianti OMS S.p.A.; the historic OMS (Officine Meccaniche) core of the Brianza "
   "district is formed, active in polyurethane machinery.")
li("2010 – Merger by incorporation of G.A.M. 85 S.r.l. (Verano Brianza).")
li("2017 – Real-estate demerger with the formation of LO.MA. Immobiliare S.r.l. (separation of the property "
   "assets).")
li("2018 – Merger by incorporation of OMS Automation S.r.l. and Polyusus Italia III S.p.A. (a Hennecke group "
   "vehicle); simultaneous change of name from \"Impianti OMS S.p.A.\" to \"HENNECKE-OMS S.P.A.\" (August 2018): "
   "entry into the orbit of the German group.")
li("2020 – Hennecke GmbH becomes sole shareholder (15/12/2020); merger by incorporation of Hennecke Italy "
   "S.r.l. Completion of the integration into the group.")
li("2026 – Renewal of the Board of Directors: departure of A.E.C. Mariani (OMS family), confirmation of "
   "R. Trippler (Chairman) and T. Wildt (MD), arrival of A.W. Mari as Managing Director.")

# ---------- 3. Products ----------
h1("3. Products, Technology & Commercial Positioning")
p("HENNECKE-OMS operates in the segment of high-tech capital goods, offering \"turnkey\" machines and plants "
  "for polyurethane processing. Its positioning is that of a process-technology supplier to industrial "
  "customers (manufacturers of foams, panels, components), with an offering that integrates design, "
  "manufacturing, installation, commissioning and after-sales service on a worldwide scale, benefiting from "
  "the commercial network of the Hennecke group.")
h2("3.1 Main product lines")
for x in [
    "High- and low-pressure metering and mixing machines and mixing heads for polyurethane",
    "Lines for flexible slabstock foam and moulded foam",
    "Continuous and discontinuous lines for insulating sandwich panels (construction, cold rooms)",
    "Plants for domestic and commercial refrigeration (PU injection)",
    "Solutions for automotive, footwear (soles), elastomers, RIM and composite materials",
    "Service, spare parts, retrofit and modernisation of the installed base (aftermarket)",
]:
    li(x)
p("The product-range detail is reconstructed from public industry sources and the group website; the "
  "breakdown of revenue by product line/geography is not available in the official documents analysed.")
h2("3.2 Technology drivers and commercial levers")
for x in [
    "Blowing-agent transition (from HFC to HFO/pentane) driving renewal of customers' machine fleet",
    "Growing automation and digitalisation of the lines (efficiency, quality, traceability)",
    "Aftermarket as a source of recurring, high-margin revenue (spare parts, service, retrofit)",
    "Cross-selling and geographic coverage through the global Hennecke platform",
]:
    li(x)

# ---------- 4. Operations ----------
h1("4. Operations & Industrial Footprint")
h2("4.1 Sites")
table(["Site", "Location", "Function"], [
    ["Headquarters & plant", "Verano Brianza (MB) – Via Lombardia / Via Alfieri", "Manufacturing and installation of PU machines/plants"],
    ["Offices", "Verano Brianza (MB) – Via Pascoli 15", "Offices / management consulting"],
    ["Closed units", "Novedrate (CO), Abbiategrasso (MI), Isola Vicentina (VI)", "Warehouse/offices closed (2010-2025)"],
])
h2("4.2 Certifications and quality")
for x in [
    "UNI EN ISO 9001:2015 – Quality Management System (TÜV Italia, first issued 2011)",
    "UNI EN ISO 14001:2015 – Environmental Management System (TÜV Italia, first issued 2024)",
    "Certified activities: design, manufacture, commissioning and service of polyurethane machines and plants",
]:
    li(x)
h2("4.3 Workforce")
p("2025 composition (source INPS): white-collar ~54%, blue-collar ~37%, middle managers ~4%, executives "
  "~4%, apprentices ~1%; 96-97% of staff are on permanent contracts and about 97% full-time. The profile "
  "reflects a project-based engineering and manufacturing company, with a predominance of technical/clerical "
  "staff. The slight headcount contraction (170 → 159 over three years) is consistent with the "
  "post-integration rationalisation phase and the 2024 volume decline.")
table(["Year", "Employees (annual average)"], [
    ["2023", "170"],
    ["2024", "166"],
    ["2025", "159"],
])

# ---------- 5. Market ----------
h1("5. Market Analysis – Polyurethane Machinery")
p("The company is classified in the sector of manufacturing machinery for the plastics and rubber industry "
  "(ATECO 28.96.00; Coface sector \"Mechanics and precision\"). It is a capital-goods segment, by nature "
  "cyclical and correlated to the capex of industrial customers. Italy and Germany represent the two "
  "historic poles of polyurethane machinery worldwide.")
h2("5.1 Demand drivers")
for x in [
    "Building insulation and energy efficiency (sandwich panels, rigid foams): EU regulatory push (EPBD, "
    "Green Deal) and renovations",
    "Domestic/commercial refrigeration and cold chain (PU injection in appliances and cold rooms)",
    "Automotive and e-mobility (lightweighting, seating, components) and furniture/mattresses (flexible foam)",
    "Footwear, elastomers and composites",
]:
    li(x)
h2("5.2 Structural trends")
for x in [
    "Decarbonisation and building regulation → demand for high-performance insulating solutions",
    "Replacement of blowing agents (HFO/pentane) → plant renewal cycle",
    "Automation, digitalisation and services (service/aftermarket)",
    "Reshoring/nearshoring of manufacturing and investment in new capacity in emerging markets",
]:
    li(x)
p("The market considerations are qualitative and based on public industry sources; no quantitative market "
  "data specific to the segment is available in the documents analysed.")

# ---------- 6. Competitor Mapping ----------
h1("6. Competitor Mapping")
p("HENNECKE-OMS competes in the global market for polyurethane process technologies, where it faces a "
  "limited number of specialised manufacturers, predominantly Italian and German. Its membership of the "
  "Hennecke Group places it among the world leaders in the sector.")
p("Qualitative mapping: consolidated revenue data of competitors is not available in the documents analysed.")
table(["Player", "Location", "Focus", "Notes"], [
    ["HENNECKE-OMS (target)", "Verano Brianza (IT)", "PU machines/plants (flexible, panels, refrig.)", "Italian division of the Hennecke Group"],
    ["Hennecke GmbH (parent)", "Sankt Augustin (DE)", "PU process technologies", "100% parent; world leader"],
    ["Cannon Group / Afros", "Peschiera Borromeo (IT)", "PU machines, metering, composites", "Main Italian competitor"],
    ["KraussMaffei", "Munich (DE)", "Reaction Process Machinery (PU)", "ChemChina group"],
    ["Frimo Group", "Lotte (DE)", "PU tooling/systems (automotive)", ""],
    ["Kurtz Ersa", "Kreuzwertheim (DE)", "Foaming (EPS/PU)", ""],
    ["Saip", "Offanengo (IT)", "PU plants / refrigeration", ""],
])
h2("6.1 Valuation benchmark (indicative)")
p("In the capital-goods / industrial-machinery segment, market EV/EBITDA multiples typically sit in a range "
  "of 6-9x for mid-sized players, with lower values (5-7x) for small-cap businesses with thin and cyclical "
  "margins. Given HENNECKE-OMS's thin margins, a robust valuation should be based on a normalised EBITDA "
  "(stripped of cycle/project effects and intercompany components) and/or on asset-based logic, rather than "
  "on the point-in-time 2024 EBITDA.")

# ---------- 7. Financial Statement Analysis ----------
h1("7. Financial Statement Analysis 2023-2024")
p("Amounts in € thousand. Source: Coface Full Report (filed financial statements 2023 and 2024). \"Coface\" "
  "EBITDA and the gross operating margin (MOL) may differ due to the different treatment of certain items "
  "(other charges, adjustments).")
h2("7.1 Reclassified Income Statement")
table(["Item", "2023", "2024", "Δ %"], [
    ["Revenue (sales & services)", "50.389", "44.243", "-12,2%"],
    ["Change in work in progress (WIP)", "+19.303", "-536", "n.m."],
    ["Other income", "1.011", "543", "-46,3%"],
    ["Value of production", "51.998", "40.652", "-21,8%"],
    ["Raw material purchases", "25.381", "14.095", "-44,5%"],
    ["Change in raw material inventory", "-1.476", "+1.523", "n.m."],
    ["Services", "10.354", "9.267", "-10,5%"],
    ["Lease/rental costs", "1.965", "1.725", "-12,2%"],
    ["Personnel costs", "11.990", "12.317", "+2,7%"],
    ["Other operating costs", "160", "183", "+14,4%"],
    ["Gross operating margin (MOL)", "2.773", "1.182", "-57,4%"],
    ["EBITDA", "3.624", "1.542", "-57,5%"],
    ["EBITDA margin", "7,2%", "3,5%", "-3,7 p.p."],
    ["Depreciation, amortisation & write-downs", "2.609", "2.549", "-2,3%"],
    ["EBIT", "1.015", "-1.007", "n.m."],
    ["Net financial income/(expense)", "-968", "-365", "+62,3%"],
    ["Adjustments / extraordinary income", "451", "1.131", "+150,8%"],
    ["Pre-tax result", "498", "-241", "n.m."],
    ["Income taxes", "209", "0", "n.m."],
    ["Net profit (loss) for the year", "289", "-241", "n.m."],
    ["Cash flow", "2.898", "2.308", "-20,4%"],
])
h2("7.2 Reclassified Balance Sheet")
table(["Item", "2023", "2024", "Δ"], [
    ["Intangible assets", "13.348", "12.016", "-1.332"],
    ["Tangible assets", "4.992", "4.148", "-844"],
    ["Financial assets", "0", "0", "0"],
    ["Net fixed assets", "18.340", "16.164", "-2.176"],
    ["Inventories", "16.706", "8.088", "-8.618"],
    ["of which work in progress", "5.940", "3.369", "-2.571"],
    ["Receivables (total)", "4.508", "4.211", "-297"],
    ["of which trade receivables", "425", "673", "+248"],
    ["Cash and cash equivalents", "8.371", "10.365", "+1.994"],
    ["Accrued income & prepaid expenses", "170", "220", "+50"],
    ["TOTAL ASSETS", "48.095", "39.048", "-9.047"],
    ["Shareholders' equity", "6.137", "6.043", "-94"],
    ["of which share capital", "2.322", "2.322", "0"],
    ["Provisions for risks and charges", "3.372", "2.758", "-614"],
    ["Employee severance (TFR)", "315", "267", "-48"],
    ["Total payables/debt", "37.687", "29.457", "-8.230"],
    ["of which customer advances", "14.319", "10.948", "-3.371"],
    ["of which trade payables", "8.096", "5.988", "-2.108"],
    ["of which medium/long-term debt", "11.578", "8.777", "-2.801"],
    ["of which bank debt (M/L term)", "0", "6.907", "+6.907"],
    ["Accrued expenses & deferred income", "584", "523", "-61"],
    ["TOTAL LIABILITIES & EQUITY", "48.095", "39.048", "-9.047"],
])
h2("7.3 Economic-financial ratios")
table(["Ratio", "2023", "2024", "Comment"], [
    ["EBITDA margin", "7,2%", "3,5%", "Compression (cycle/project effect)"],
    ["MOL margin", "5,5%", "2,7%", "Declining"],
    ["EBIT margin", "2,0%", "-2,3%", "Negative operating result in 2024"],
    ["Net margin", "0,6%", "-0,5%", "Net loss in 2024"],
    ["Quick ratio (liquidity)", "48,9", "69,8", "Improved; ample liquidity"],
    ["Debt ratio", "60,1", "50,1", "Decreasing"],
    ["Equity / Total assets", "12,8%", "15,5%", "Thin equity (subsidiary)"],
    ["DSO trade (days)", "3,0", "5,5", "Sales with advances (minimal receivables)"],
    ["DIO inventory (days)", "119,4", "65,8", "Strong inventory reduction"],
    ["DPO suppliers (days)", "85,1", "86,6", "Stable"],
    ["Net working capital", "3.476", "1.984", "Reduced (€ thousand)"],
])
h2("7.4 Concise reading")
for x in [
    "Declining revenue (-12%) and a sharp drop in value of production (-22%): the differential is entirely "
    "explained by the change in work in progress, typical of the long project-cycle model (strong WIP "
    "build-up in 2023, run-down in 2024).",
    "Compressed margins: EBITDA from €3.6M to €1.5M (margin from 7.2% to 3.5%), with a negative operating "
    "result (-€1.0M) and a net loss of €0.24M. Personnel cost, essentially stable (+2.7%), weighed in "
    "relative terms on the lower value of production.",
    "Structural reduction of variable costs: raw material purchases -44.5% (-€11.3M) and services -10.5%, "
    "consistent with the lower production load.",
    "Sharp reduction of working capital and inventory: inventories from €16.7M to €8.1M (DIO from 119 to 66 "
    "days); net working capital falls to €2.0M. The model is self-financed by customer advances (€10.9M).",
    "Solid financial position in terms of liquidity: cash of €10.4M against medium/long-term bank debt of "
    "€6.9M (new in 2024); total debt (€29.5M) consists largely of operating items (advances, suppliers) and "
    "intercompany positions.",
    "Thin equity (€6.0M) and high intangible assets (€12.0M): the balance-sheet structure is typical of a "
    "subsidiary, with goodwill/intangibles from combinations to be monitored (impairment risk).",
]:
    li(x)

# ---------- 8. Earnings dynamics ----------
h1("8. Earnings Dynamics & Variance Analysis")
p("The fall in EBITDA from 2023 to 2024 (from €3.6M to €1.5M; MOL from €2.8M to €1.2M) must be read in light "
  "of the project-based nature of the business. In 2023 the company had strongly increased work in progress "
  "(+€19.3M), inflating the value of production to €52.0M; in 2024 the WIP change turned negative (-€0.5M) "
  "and the value of production fell to €40.7M. Absolute margin therefore shrank mainly due to the lower "
  "absorption of fixed costs (chiefly personnel, stable in value but rising in incidence) over a lower "
  "production volume.")
h2("8.1 Main variances (€ thousand, 2024 vs 2023)")
table(["Item", "Change", "Effect on EBITDA"], [
    ["Value of production", "-11.346", "Unfavorable"],
    ["Raw material purchases (lower cost)", "-11.286", "Favorable"],
    ["Change in raw material inventory", "+2.999", "Favorable"],
    ["Services (lower cost)", "-1.087", "Favorable"],
    ["Lease/rental (lower cost)", "-240", "Favorable"],
    ["Personnel cost (higher cost)", "+327", "Unfavorable"],
    ["Other operating costs", "+23", "Unfavorable"],
    ["Net change in gross operating margin", "-1.591", ""],
])
p("Financial management, on the other hand, improved: net financial expense from -€0.97M to -€0.37M, and "
  "extraordinary items/adjustments contributed positively by €1.1M (vs €0.5M), reducing the impact of the "
  "operating loss on the pre-tax result (-€0.24M). In the absence of taxes (loss), the net result coincides "
  "with the pre-tax result.")
h2("8.2 Implications for normalisation")
for x in [
    "The point-in-time 2024 EBITDA (3.5%) is depressed by the cycle: a normalisation over several years "
    "(average and backlog) is essential to estimate the sustainable run-rate.",
    "Intercompany components (transfer prices on purchases/sales and service with Hennecke GmbH) must be "
    "isolated to measure \"stand-alone\" profitability.",
    "Extraordinary items/adjustments (€1.1M in 2024) require an analysis of recurrence.",
]:
    li(x)

# ---------- 9. PE considerations ----------
h1("9. Considerations for Private Equity")
p("HENNECKE-OMS presents an atypical profile relative to a classic PE target: it is a captive subsidiary "
  "with thin margins, integrated into an industrial group. The following considerations apply both in the "
  "scenario of a carve-out of the Italian entity and in that of an acquisition of the entire Hennecke Group.")
h2("9.1 Points of attractiveness")
for x in [
    "Historic brand (OMS) and distinctive know-how on PU plants (sandwich panels, flexible foam), with a "
    "global installed base that feeds the aftermarket.",
    "Exposure to structural drivers (building energy efficiency, cold chain, e-mobility, blowing-agent "
    "transition).",
    "Light financial structure: low bank debt, ample liquidity, working capital self-financed by advances.",
    "Room to recover margins (from 3.5% towards 6-8%) via fixed-cost efficiency and a service/spare-parts mix.",
]:
    li(x)
h2("9.2 Critical issues and cautions")
for x in [
    "Thin and volatile margins; negative operating result in 2024.",
    "Captive nature: limited managerial, commercial and treasury autonomy; dependence on intercompany "
    "relationships (complex carve-out: Hennecke brand, IT, shared services).",
    "Reduced equity (€6.0M) and high intangible assets (€12.0M) to be tested.",
    "Pledged shares; need to release the constraint in case of a sale.",
    "Cyclicality and execution risk on complex projects.",
]:
    li(x)
h2("9.3 Value-creation levers")
for x in [
    "Backlog normalisation and pricing discipline on projects.",
    "Development of the aftermarket (service, spare parts, retrofit) as an engine of recurring revenue and "
    "margin.",
    "Industrial efficiency and rationalisation of fixed costs.",
    "Working-capital optimisation (already under way: DIO from 119 to 66 days).",
    "Possible buy-and-build in PU/adjacent capital-goods machinery.",
]:
    li(x)

# ---------- 10. DD ----------
h1("10. Areas for Further Due Diligence")
for x in [
    "Intercompany relationships: transfer pricing on purchases/sales with Hennecke GmbH and group companies, "
    "cash pooling/treasury, any brand royalties, intercompany financing.",
    "Pledge on shares: nature, beneficiary and secured obligations; conditions for release in case of a sale.",
    "Order book (backlog) and project margins: composition, progress, expected margins, penalties and "
    "guarantees.",
    "Intangible assets (€12.0M): nature (goodwill from mergers, capitalised development), impairment testing "
    "and residual useful life.",
    "Working capital and customer advances (€10.9M): sustainability, guarantees issued, contractual "
    "milestones.",
    "Litigation and plant warranties, product liability, performance bonds.",
    "Dependence on key personnel and technical know-how; retention plan for management and technicians.",
    "Carve-out feasibility: IT systems, shared services, use of the Hennecke brand, continuity of the global "
    "sales network.",
    "EBITDA normalisation: one-offs, extraordinary and intercompany components; sustainable run-rate.",
    "\"True\" net financial position: reclassification of advances and intercompany debt versus actual "
    "financial debt.",
]:
    li(x)

# ---------- 11. SWOT ----------
h1("11. SWOT Analysis")
h2("STRENGTHS")
for x in [
    "Historic OMS brand and distinctive know-how on PU plants (continuous sandwich panels, flexible foam)",
    "Membership of the Hennecke Group, world leader: R&D, sales network and global installed base",
    "High liquidity (€10.4M) and contained bank debt (€6.9M M/L term)",
    "Working capital self-financed by customer advances (€10.9M)",
    "ISO 9001/14001 certifications; qualified workforce (54% white-collar/technical)",
    "Strong inventory reduction in 2024 (DIO from 119 to 66 days)",
]:
    li(x)
h2("WEAKNESSES")
for x in [
    "Thin and volatile margins (EBITDA 3.5% in 2024; net loss of €0.24M)",
    "Reduced equity (€6.0M; 15.5% of total assets)",
    "High intangible assets (€12.0M) from combinations: impairment risk",
    "Strong dependence on intercompany relationships and group treasury (limited autonomy)",
    "Cyclicality of revenue and execution risk on complex projects",
    "Shares pledged (constraint)",
]:
    li(x)
h2("OPPORTUNITIES")
for x in [
    "Structural drivers: building energy efficiency (insulating panels), cold chain, e-mobility",
    "Growth of aftermarket/service and spare parts (recurring high-margin revenue)",
    "Blowing-agent transition (HFO/pentane): renewal cycle of customers' machine fleet",
    "Group synergies and cross-selling; US/Asia expansion through the Hennecke network",
    "Margin recovery via backlog normalisation and fixed-cost efficiency",
]:
    li(x)
h2("THREATS")
for x in [
    "Customers' investment cycle (capital goods) and macroeconomic slowdowns",
    "Qualified competition (Cannon, KraussMaffei) and price pressure on projects",
    "FX risk on export projects and country risk",
    "Volatility of energy, steel and component costs; supply-chain tensions",
    "Concentration on individual projects and warranty/performance risk on plants",
]:
    li(x)

# ---------- 12. Sources ----------
h1("12. Sources & Disclaimer")
h2("12.1 Documents analysed")
for x in [
    "Coface Business Information – Full Report HENNECKE-OMS S.p.A. (produced on 04/06/2026): identifying "
    "data, credit score, 2023-2024 financial statements, ratios, employees, properties, certifications.",
    "Historical company extract (visura) – Chamber of Commerce of Milano-Monza-Brianza-Lodi (extracted on "
    "04/06/2026, 158 pages): company details, capital, shareholders, directors, statutory auditors, "
    "attorneys, mergers/demergers, activities, local units, history of amendments.",
]:
    li(x)
h2("12.2 Public industry sources")
for x in [
    "Official Hennecke Group website and industry materials (product range, positioning)",
    "Industry literature on polyurethane machinery and capital goods (market drivers)",
]:
    li(x)
h2("12.3 Disclaimer")
p("This document has been prepared for information purposes from public data, filed financial statements and "
  "available chamber-of-commerce documentation. The assessments expressed do not constitute an investment "
  "recommendation or professional advice; any investment decision must be preceded by specialist Due "
  "Diligence (financial, commercial, operational, legal, tax). The data have been reported to the best "
  "understanding of the sources analysed. The consolidated financial statements of the Hennecke Group "
  "(Germany) and the ownership arrangements upstream of Hennecke GmbH were not part of this analysis and may "
  "present different values and configurations. The 2024 margin reflects cycle/project effects and requires "
  "full normalisation during DD to estimate the sustainable run-rate. Product, market and competitor "
  "information is qualitative and based on public industry sources.")

# ---------- 13. Appendix ----------
h1("13. Appendix – Financial Schedules (extended format)")
p("Amounts in € thousand. Source: Coface Full Report (extended-format financial statements, Fourth EEC "
  "Directive schema).")
h2("13.1 Income Statement (extended)")
table(["Item", "2023", "2024"], [
    ["Value of production", "51.998", "40.652"],
    ["Revenue", "31.086", "44.779"],
    ["Change in work in progress", "+19.303", "-536"],
    ["Other income", "1.011", "543"],
    ["Production costs", "50.983", "41.659"],
    ["Raw material purchases", "25.381", "14.095"],
    ["Services", "10.354", "9.267"],
    ["Personnel costs", "11.990", "12.317"],
    ["Lease/rental costs", "1.965", "1.725"],
    ["Change in raw materials", "-1.476", "+1.523"],
    ["Other operating costs", "160", "183"],
    ["MOL (gross operating margin)", "2.773", "1.182"],
    ["Depreciation, amortisation & write-downs", "2.609", "2.549"],
    ["Core operating result", "164", "-1.367"],
    ["Non-core income/(expense)", "870", "705"],
    ["Financial expense", "987", "710"],
    ["Extraordinary income/(expense) (adjustments)", "451", "1.131"],
    ["Income taxes", "209", "0"],
    ["Net profit (loss) for the year", "289", "-241"],
    ["Cash flow", "2.898", "2.308"],
])
h2("13.2 Balance Sheet – Assets (extended)")
table(["Item", "2023", "2024"], [
    ["B) Fixed assets", "18.340", "16.164"],
    ["I. Intangible", "13.348", "12.016"],
    ["II. Financial", "0", "0"],
    ["III. Tangible", "4.992", "4.148"],
    ["C) Current assets", "29.585", "22.664"],
    ["I. Inventories", "16.706", "8.088"],
    ["of which work in progress", "5.940", "3.369"],
    ["II. Receivables", "4.508", "4.211"],
    ["of which trade receivables", "425", "673"],
    ["III. Cash and equivalents", "8.371", "10.365"],
    ["D) Accruals and deferrals", "170", "220"],
    ["TOTAL ASSETS", "48.095", "39.048"],
])
h2("13.3 Balance Sheet – Liabilities & Equity (extended)")
table(["Item", "2023", "2024"], [
    ["A) Shareholders' equity", "6.137", "6.043"],
    ["of which share capital", "2.322", "2.322"],
    ["B) Provisions for risks and charges", "3.372", "2.758"],
    ["C) Employee severance (TFR)", "315", "267"],
    ["D) Payables/Debt", "37.687", "29.457"],
    ["Trade payables", "8.096", "5.988"],
    ["Advances / prepayments", "14.319", "10.948"],
    ["Short-term debt", "26.109", "20.680"],
    ["Medium/long-term debt", "11.578", "8.777"],
    ["of which bank debt (M/L term)", "0", "6.907"],
    ["E) Accruals and deferrals", "584", "523"],
    ["TOTAL LIABILITIES & EQUITY", "48.095", "39.048"],
])

doc.save(OUT)
print("Saved:", OUT)
print("Paragraphs:", len(doc.paragraphs), "Tables:", len(doc.tables))
