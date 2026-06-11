#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Tailored CV (EN) for Giacomo Battaglia — bullets reordered to foreground the
CFO / Hennecke-OMS match. Facts unchanged; only ordering & framing tuned."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/home/user/Storage/Hennecke/candidacy/GBattaglia_CV_CFO_tailored_EN.docx"
BLUE = RGBColor(0x1F, 0x3A, 0x5F)
GRAY = RGBColor(0x55, 0x55, 0x55)

doc = Document()
st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(10.5)
for sec in doc.sections:
    sec.top_margin = Inches(0.5); sec.bottom_margin = Inches(0.5)
    sec.left_margin = Inches(0.7); sec.right_margin = Inches(0.7)

def hr(p):
    pPr = p._p.get_or_add_pPr(); b = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6'); bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F3A5F'); b.append(bottom); pPr.append(b)

def section(title):
    p = doc.add_paragraph(); p.space_before = Pt(8); p.space_after = Pt(2)
    r = p.add_run(title.upper()); r.bold = True; r.font.size = Pt(11.5); r.font.color.rgb = BLUE
    hr(p)
    return p

def bullet(runs, after=2):
    p = doc.add_paragraph(style="List Bullet"); p.space_after = Pt(after)
    p.paragraph_format.left_indent = Inches(0.2); p.paragraph_format.first_line_indent = Inches(-0.12)
    if isinstance(runs, str): runs = [(runs, False)]
    for t, b in runs:
        r = p.add_run(t); r.bold = b; r.font.size = Pt(10)
    return p

def line(runs, size=10.5, after=1, align=None, space_before=0):
    p = doc.add_paragraph(); p.space_after = Pt(after); p.space_before = Pt(space_before)
    if align: p.alignment = align
    if isinstance(runs, str): runs = [(runs, False, None)]
    for tup in runs:
        t = tup[0]; b = tup[1] if len(tup) > 1 else False
        col = tup[2] if len(tup) > 2 else None
        r = p.add_run(t); r.bold = b; r.font.size = Pt(size)
        if col is not None: r.font.color.rgb = col
    return p

# ---------------- Header ----------------
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(0)
r = p.add_run("GIACOMO BATTAGLIA"); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = BLUE
line([("Group CFO  |  C-Level & Board Member  |  M&A & Value Creation  |  FP&A Director", True)],
     size=11, after=1, align=WD_ALIGN_PARAGRAPH.CENTER)
line([("Target role: Group / Division CFO — industrial manufacturing within a PE / industrial-group structure", False, GRAY)],
     size=9.5, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
line([("Milan, Italy  ·  Italian (native), English (fluent)  ·  MSc Economics & Corporate Management", False, GRAY)],
     size=9.5, after=1, align=WD_ALIGN_PARAGRAPH.CENTER)
line([("giacomo.battaglia80@gmail.com  ·  +39 349 7908205  ·  linkedin.com/in/giacomobattaglia", False, GRAY)],
     size=9.5, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)

# ---------------- Executive summary ----------------
section("Executive Summary")
line("Group CFO with 15+ years in multinational and private-equity-backed industrial environments "
     "(Manufacturing, Chemistry, FMCG), with full P&L, cash-flow and balance-sheet ownership of businesses up to "
     "€400m turnover. Recognised value-creation partner to PE funds, shareholders and foreign parent companies, "
     "operating fluently in matrix and centralised-governance settings (dotted-line reporting to UK / Swiss group "
     "CFOs). Deep expertise in industrial controlling and margin recovery, working-capital and cash optimisation, "
     "transfer pricing and intercompany governance, M&A and carve-outs, and SAP S/4HANA finance transformation.",
     size=10.5, after=3)

# ---------------- Core strengths (mapped to role) ----------------
section("Core Strengths Most Relevant to the Role")
bullet([("Industrial controlling & margin recovery — ", True),
        ("standard costing, variance analysis (Price/Mix/PPV), OEE, scrap, give-away & mass balance, pricing-waterfall governance.", False)])
bullet([("Working capital & cash optimisation — ", True),
        ("non-recourse factoring (DSO), reverse factoring (DPO), VMI/VMR (DIO), sale-and-leaseback, cross-border cash pooling.", False)])
bullet([("Transfer pricing, tax & intercompany governance — ", True),
        ("renegotiated a €30M tax assessment down to a €5M settlement; intercompany pricing & financing across multiple entities.", False)])
bullet([("M&A, carve-outs & PE value creation — ", True),
        (">€100m cumulative deals (equity & asset), post-merger integration, due diligence; active participant in a PE exit (CapVest → Bain Capital).", False)])
bullet([("Finance transformation & reporting — ", True),
        ("SAP ECC → S/4HANA, Power BI / OneStream frameworks; dual statutory reporting (OIC + IAS/IFRS); former KPMG auditor.", False)])

# ---------------- Selected achievements ----------------
section("Selected Achievements & Projects")

line([("Transfer Pricing & Tax", True, BLUE)], size=10.5, after=1, space_before=2)
bullet([("Negotiated a win-win settlement with the Italian Tax Authority, reducing a ", False), ("€30M assessment to a €5M settlement", True),
        (", by challenging the group's management-fee recharge methodology and the deductibility of interest on the acquisition master financing (CapVest fund / Balconi). Partner: PwC.", False)])

line([("Cash Flow & Working Capital Optimisation", True, BLUE)], size=10.5, after=1, space_before=3)
bullet("Non-recourse factoring to improve DSO (Factor France, Unicredit Factoring, FinDynamic, BPER).")
bullet("Reverse factoring to extend DPO (Crédit Agricole Factoring, Banco BPM); VMI/VMR to optimise DIO.")
bullet("Sale-and-leaseback on owned buildings to strengthen liquidity and fund capacity expansion (Ascension UK); multi-bank cash pooling (Citibank, Barclays).")

line([("M&A, Integration & Carve-outs (>€100m cumulative)", True, BLUE)], size=10.5, after=1, space_before=3)
bullet([("Equity acquisition & merger of ", False), ("Idp Srl", True), (" — debt restructuring and bank guarantee pledging shares & the Balconi brand (Vitale, NDA, LCA, Pirola, Crédit Agricole / Dentons / Baker McKenzie).", False)])
bullet([("Equity acquisition of ", False), ("Freddi S.p.A.", True), ("; merger of ", False), ("Dolciaria Valdenza", True), (" (incl. ERP integration).", False)])
bullet([("Asset deals (court auctions): ", False), ("Dal Colle", True), (" (Verona), ", False), ("Melegatti", True), (" (incl. asset appraisal), ", False), ("Berkel by Rovagnati", True), (" (Milan).", False)])
bullet("Set-up of two French subsidiaries (a holding to support future M&A and a commercial entity for direct key-account management); coordination of greenfield international subsidiaries.")

line([("Capex & Investment Appraisal", True, BLUE)], size=10.5, after=1, space_before=3)
bullet([("€80m", True), (" of capacity investment appraised and financed over three years (two new wafer lines, one sponge-cake line) via business case, NPV, IRR, payback and scenario analysis.", False)])

line([("Pricing, Planning & Board Support", True, BLUE)], size=10.5, after=1, space_before=3)
bullet("Reshaped the global pricing waterfall to support commercial excellence (Adama global pricing / Wiglaf); EU lead for the 5-year strategic & financial plan in Power BI.")
bullet("Monthly Board review of financial/economic KPIs across operations, logistics and procurement; variance analysis on Price/Mix/PPV/synergies/productivity.")

line([("ERP & Digital Finance Transformation", True, BLUE)], size=10.5, after=1, space_before=3)
bullet("Sponsor of SAP ECC → S/4HANA migration (completed Dec 2025, live Jan 2026; partner Altea Up) and Power BI reporting framework; hands-on with QlikView, OneStream, GFS, Power Automate, O365; building AI agents (Claude Code / Copilot) to automate recurring analysis.")

# ---------------- Professional experience ----------------
section("Professional Experience")

def role(years, title, company, meta):
    p = doc.add_paragraph(); p.space_after = Pt(0); p.space_before = Pt(4)
    r = p.add_run(title + "  —  " + company); r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = BLUE
    r2 = p.add_run("    " + years); r2.italic = True; r2.font.size = Pt(9.5); r2.font.color.rgb = GRAY
    line([(meta, False, GRAY)], size=9.5, after=1)

role("2020 – Present", "Group Chief Financial Officer & Board Member", "Valeo South & West Europe (Food / FMCG · Bain Capital portfolio)",
     "CFO of 4 Italian entities (Valeo Bidco Italia / Balconi / Idp / Freddi) · €350m turnover (50% export) · 8 sites · dotted line to UK Group CFO · finance team of 25 FTE")
bullet("Full P&L, cash-flow and balance-sheet accountability; Deputy CEO on commercial strategy & organisation; leadership of Finance, Commercial/Industrial Controlling, Treasury, Tax & Transfer Price, IT and Internal Audit.")
bullet("Lead role in multiple M&A transactions (>€100m cumulative PPA); active participant in the PE exit from CapVest to Bain Capital.")
bullet("Cash freed through working-capital and treasury optimisation; €80m capex appraised and financed across three years.")
bullet("Sponsor of SAP ECC → S/4HANA migration; dual statutory reporting (OIC) and IAS/IFRS.")

role("2017 – 2020", "Chief Financial Officer Italy", "Adama Italia (Chemistry / Crop Protection · Syngenta Group – ChemChina)",
     "€100m local · $4bn global · dotted line to EMEA Group CFO (Switzerland) · Finance/HR team of 8 FTE")
bullet("Owner of pricing architecture and margin governance; five-year strategic planning & Power BI reporting framework.")
bullet("SOX, compliance, corporate governance & GDPR lead; leadership of Finance, Commercial, Treasury, Tax & Transfer Price, IT/HR; dual OIC and IAS/IFRS reporting.")

role("2015 – 2017", "Finance Manager", "Berkel S.r.l. (Industrial appliances · Rovagnati Group)",
     "€30m · sites in Italy + India (greenfield), US, Germany · reporting to Board and Group CFO")
bullet("Full P&L, cash flow, budgeting and forecasting; industrial controlling and margin governance (standard costing, variance analysis, transfer pricing); coordination of international subsidiaries.")

role("2005 – 2015", "Earlier roles (controlling, FP&A & audit)", "",
     "Head of FP&A — Thule Italia (Automotive, 2014–15) · FP&A Senior — Vodafone Italia (Telco, 2011–14) · FP&A Senior — Beta (tools manufacturing, 2008–11) · Internal Audit — Saipem (Oil & Gas, 2007–08) · Audit — KPMG (2005–07)")

# ---------------- Skills & tools ----------------
section("Skills, Tools & Education")
bullet([("Finance & control: ", True), ("M&A / PMI / due diligence, value creation, working capital & treasury, transfer pricing & tax, industrial & commercial controlling, FP&A, 5-year planning, Board reporting.", False)])
bullet([("Systems: ", True), ("SAP S/4HANA & ECC, OneStream, Power BI, QlikView, GFS, Power Automate, Microsoft 365; AI agents (Claude Code / Copilot).", False)])
bullet([("Education & languages: ", True), ("MSc in Economics & Corporate Management · Italian (native), English (fluent).", False)])

doc.save(OUT)
print("Saved:", OUT, "| paragraphs:", len(doc.paragraphs))
