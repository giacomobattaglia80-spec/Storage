#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Study guide: Project / contract accounting (WIP, % completion, IFRS 15 & OIC 23)
tailored to HENNECKE-OMS, for interview preparation."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/home/user/Storage/Hennecke/candidacy/Study_Guide_Project_Accounting_IFRS15.docx"
BLUE = RGBColor(0x1F, 0x3A, 0x5F); RED = RGBColor(0xC8, 0x10, 0x2E); GRAY = RGBColor(0x55, 0x55, 0x55)

doc = Document()
s = doc.styles["Normal"]; s.font.name = "Calibri"; s.font.size = Pt(10.5)
for sec in doc.sections:
    sec.top_margin = Inches(0.6); sec.bottom_margin = Inches(0.6)
    sec.left_margin = Inches(0.75); sec.right_margin = Inches(0.75)

def hr(p):
    pPr = p._p.get_or_add_pPr(); b = OxmlElement('w:pBdr'); bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6'); bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), '1F3A5F')
    b.append(bot); pPr.append(b)

def h1(t):
    p = doc.add_paragraph(); p.space_before = Pt(10); p.space_after = Pt(2)
    r = p.add_run(t); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = BLUE; hr(p)

def h2(t):
    p = doc.add_paragraph(); p.space_before = Pt(6); p.space_after = Pt(1)
    r = p.add_run(t); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = BLUE

def para(runs, size=10.5, after=3, before=0, align=None, color=None):
    p = doc.add_paragraph(); p.space_after = Pt(after); p.space_before = Pt(before)
    if align: p.alignment = align
    if isinstance(runs, str): runs = [(runs, False)]
    for tup in runs:
        t = tup[0]; b = tup[1] if len(tup) > 1 else False; c = tup[2] if len(tup) > 2 else color
        r = p.add_run(t); r.bold = b; r.font.size = Pt(size)
        if c is not None: r.font.color.rgb = c
    return p

def bullet(runs, after=2):
    p = doc.add_paragraph(style="List Bullet"); p.space_after = Pt(after)
    if isinstance(runs, str): runs = [(runs, False)]
    for tup in runs:
        t = tup[0]; b = tup[1] if len(tup) > 1 else False
        r = p.add_run(t); r.bold = b; r.font.size = Pt(10)

def num(runs, after=2):
    p = doc.add_paragraph(style="List Number"); p.space_after = Pt(after)
    if isinstance(runs, str): runs = [(runs, False)]
    for tup in runs:
        t = tup[0]; b = tup[1] if len(tup) > 1 else False
        r = p.add_run(t); r.bold = b; r.font.size = Pt(10)

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for i, hdr in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        rr = c.paragraphs[0].add_run(hdr); rr.bold = True; rr.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""; rr = cells[i].paragraphs[0].add_run(str(v)); rr.font.size = Pt(9.5)
    return t

# ===== Title =====
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(0)
r = p.add_run("Project / Contract Accounting — Study Guide"); r.bold = True; r.font.size = Pt(17); r.font.color.rgb = BLUE
para([("WIP · percentage-of-completion · IFRS 15 & OIC 23 — tailored to HENNECKE-OMS (interview prep)", False, GRAY)],
     size=10.5, after=6, align=WD_ALIGN_PARAGRAPH.CENTER)

# ===== 0. Why it matters =====
h1("0. Perché conta (e come spiega il loro bilancio)")
para([("Hennecke-OMS è un business ", False), ("engineer-to-order / a commessa lunga", True),
      (": progetta e costruisce impianti su misura in 12–24 mesi. La contabilità di commessa è il motivo per cui, nel dossier, ", False),
      ("ricavi e “valore della produzione” divergono", True), (" e per cui la marginalità oscilla. Saperlo spiegare è il tuo biglietto da visita tecnico.", False)])
para([("Numeri loro da saper leggere a memoria: ", True),
      ("la variazione lavori in corso (WIP) è passata da ", False), ("+€19,3M (2023) a -€0,5M (2024)", True),
      ("; per questo il valore della produzione è sceso -22% mentre i ricavi solo -12%. Gli ", False),
      ("acconti da clienti erano €10,9M", True), (" e i lavori in corso a magazzino €3,4M (2024).", False)])
para([("Messaggio da colloquio: ", True),
      ("“Il crollo 2024 dell'EBITDA è in larga parte un effetto di timing della commessa (smaltimento del WIP accumulato nel 2023), non un deterioramento strutturale. Per questo la valutazione va fatta su un EBITDA normalizzato e su backlog/margine-a-finire, non sul dato puntuale.”", False, RED)])

# ===== 1. Two frameworks =====
h1("1. I due mondi che gestirai: OIC 23 (statutory) e IFRS 15 (gruppo)")
para([("Hennecke-OMS tiene ", False), ("doppia contabilità", True),
      (": bilancio civilistico italiano (OIC) e reporting di gruppo (IAS/IFRS). Le commesse lunghe sono trattate da:", False)])
bullet([("OIC 23 – “Lavori in corso su ordinazione”", True), (" (Italian GAAP).", False)])
bullet([("IFRS 15 – “Revenue from Contracts with Customers”", True), (" (gruppo).", False)])
para([("Concetto chiave comune: ", True), ("quando i criteri sono soddisfatti, ricavi e margine si riconoscono ", False),
      ("nel tempo (over time / percentuale di completamento)", True), (", non alla consegna. Quando non lo sono, si riconoscono ", False),
      ("a un punto nel tempo (alla consegna/collaudo)", True), (".", False)])

# ===== 2. IFRS 15 five steps =====
h1("2. IFRS 15 — il modello in 5 step")
num([("Identifica il contratto", True), (" con il cliente (diritti/obblighi esecutivi, incasso probabile).", False)])
num([("Identifica le performance obligations (PO)", True), (" — le prestazioni distinte: es. macchina, installazione, commissioning, training, ricambi, service/garanzia estesa.", False)])
num([("Determina il prezzo della transazione", True), (" — incluso il ", False), ("variable consideration", True), (" (penali, bonus, milestone) e l'eventuale ", False), ("componente finanziaria significativa", True), (" sugli acconti.", False)])
num([("Alloca il prezzo", True), (" alle PO in base allo ", False), ("stand-alone selling price", True), (".", False)])
num([("Riconosci il ricavo", True), (" quando/man mano che la PO è soddisfatta — ", False), ("over time", True), (" oppure ", False), ("point in time", True), (".", False)])

# ===== 3. Over time vs point in time =====
h1("3. Over time o point in time? (il giudizio chiave per le macchine)")
para([("Si riconosce ", False), ("OVER TIME", True), (" se è soddisfatto almeno uno dei 3 criteri (IFRS 15.35):", False)])
bullet([("(a) ", True), ("il cliente riceve e consuma i benefici man mano (tipico dei servizi continuativi);", False)])
bullet([("(b) ", True), ("la prestazione crea/migliora un asset che il cliente controlla mentre è costruito;", False)])
bullet([("(c) ", True), ("l'asset ", False), ("non ha uso alternativo", True), (" per il fornitore (è su misura) ", False), ("e", True), (" c'è ", False), ("diritto contrattuale al pagamento per il lavoro svolto a oggi", True), (" (costi + margine ragionevole).", False)])
para([("Per impianti su misura, spesso vale il ", False), ("(c)", True),
      (" → riconoscimento over time. Se invece l'asset è standard/rivendibile o non c'è diritto al pagamento progressivo → ", False),
      ("point in time", True), (" alla consegna/collaudo. È una ", False), ("scelta contrattuale", True),
      (": clausole su titolo, accettazione e diritto al pagamento determinano il trattamento. Da verificare in DD.", False)])

# ===== 4. Measuring progress =====
h1("4. Come si misura l'avanzamento (% completion)")
para([("Input method – cost-to-cost (il più comune):", True)])
para([("   % avanzamento = costi sostenuti a oggi ÷ costi totali stimati (EAC).  Ricavo del periodo = % × prezzo della transazione.", False)], size=10)
para([("Output method:", True), (" milestone, unità prodotte, perizie tecniche di avanzamento fisico.", False)])
para([("Attenzione: ", True), ("i materiali a piè d'opera non ancora “consumati” e gli ", False),
      ("uplift iniziali (mobilization)", True), (" possono distorcere il cost-to-cost → si rettificano. La qualità della stima ", False),
      ("cost-to-complete", True), (" è il cuore del controllo di commessa.", False)])

# ===== 5. Contract asset/liability =====
h1("5. Contract asset vs contract liability (la meccanica del WIP e degli acconti)")
table(["Posta", "Cos'è", "Dove (IFRS 15)", "Nel caso Hennecke-OMS"],
      [["Contract asset", "Ricavo maturato ma non ancora fatturato (over-billing del lavoro sul cliente)", "Attivo", "Lavori in corso su ordinazione (€3,4M, 2024)"],
       ["Receivable", "Credito incondizionato (già fatturato, manca solo l'incasso)", "Attivo (crediti v/clienti)", "Crediti v/clienti (bassi: €0,7M)"],
       ["Contract liability", "Fatturato/acconti incassati in eccesso rispetto al ricavo maturato", "Passivo", "Acconti da clienti (€10,9M)"]])
para([("Lettura: ", True), ("acconti elevati + WIP contenuto → il circolante è ", False),
      ("autofinanziato dai clienti", True), (" (DSO bassissimo, ~3–5 gg). È una forza del modello, ma gli acconti sono un ", False),
      ("debito/obbligazione di performance", True), (", non cassa “libera”.", False)])

# ===== 6. OIC 23 =====
h1("6. OIC 23 — “Lavori in corso su ordinazione” (lato civilistico)")
bullet([("Se i ricavi e lo stato avanzamento sono stimabili in modo attendibile → ", False), ("criterio della percentuale di completamento", True), (" (di norma cost-to-cost): ricavi e margine maturano nel tempo.", False)])
bullet([("Altrimenti → ", False), ("criterio della commessa completata", True), (" (ricavo a fine commessa).", False)])
bullet([("La ", False), ("variazione dei lavori in corso su ordinazione", True), (" transita a conto economico nel ", False), ("Valore della produzione (voce A.3)", True), (" → ", False), ("ecco perché il loro “valore della produzione” gonfia/sgonfia con il WIP.", True)])
bullet([("Perdite previste a finire → ", False), ("fondo per rischi su commesse (onerous contract)", True), (" rilevato subito.", False)])

# ===== 7. Worked example =====
h1("7. Esempio numerico (da saper rifare alla lavagna)")
para([("Commessa: macchina su misura, prezzo €10,0M, costi totali stimati €8,0M (margine atteso €2,0M, 20%). Durata 2 anni. Metodo cost-to-cost. Acconto €3,0M all'ordine.", False)])
table(["", "Anno 1", "Anno 2", "Totale"],
      [["Costi sostenuti nel periodo", "6,0", "2,0", "8,0"],
       ["% avanzamento cumulata", "75%", "100%", "—"],
       ["Ricavo riconosciuto (10,0 × %)", "7,5", "2,5", "10,0"],
       ["Margine riconosciuto", "1,5", "0,5", "2,0"],
       ["Variazione WIP / contract asset", "+1,5*", "−1,5", "0"]])
para([("* In Anno 1: ricavo 7,5 vs fatturato/acconto 3,0 → ", False), ("contract asset (WIP) = 4,5", True),
      (" al netto si riflette nella “variazione lavori in corso”; in Anno 2 il WIP si smaltisce (segno negativo). ", False),
      ("Questo è esattamente il meccanismo dietro il loro swing +€19,3M → -€0,5M.", True)])
para([("Scrittura tipica Anno 1 (semplificata): ", True),
      ("Costi a fornitori/magazzino; a fine periodo si rileva il ricavo maturato (7,5) con contropartita contract asset (al netto degli acconti girati a contract liability).", False)], size=10)

# ===== 8. Judgements & risks =====
h1("8. Stime, giudizi e rischi (dove un CFO crea o brucia valore)")
bullet([("Cost-to-complete & EAC", True), (": revisione periodica; un errore di stima ribalta il margine (“margin-at-completion”).", False)])
bullet([("Variable consideration", True), (": penali per ritardo, bonus, claim — da stimare con il “constraint” (solo se altamente probabile).", False)])
bullet([("Onerous contracts", True), (": perdita attesa → accantonamento immediato e integrale.", False)])
bullet([("Contract modifications / varianti", True), (": se distinte → nuovo contratto; altrimenti catch-up cumulativo.", False)])
bullet([("Significant financing component", True), (": acconti rilevanti (€10,9M) possono contenere una componente finanziaria implicita.", False)])
bullet([("Warranty & performance bond", True), (": garanzia di legge (costo) vs garanzia-servizio (PO separata); fideiussioni e penali.", False)])
bullet([("Intercompany", True), (": prezzi di trasferimento su forniture/ingegneria con Hennecke GmbH → isolare il margine stand-alone (collega al transfer pricing, tuo punto forte).", False)])

# ===== 9. KPI =====
h1("9. KPI e “red flags” di controllo commessa")
bullet([("Backlog / book-to-bill", True), (", margine-a-finire per commessa, EAC vs budget.", False)])
bullet([("Over/under-billing", True), (" (contract asset vs liability) e cash-curve della commessa (milestone di incasso).", False)])
bullet([("Slittamenti costi/tempi", True), (", indice di affidabilità delle stime, perdite previste, claim aperti.", False)])
bullet([("Normalizzazione EBITDA", True), (": depurare swing di WIP, one-off e infragruppo → run-rate sostenibile.", False)])

# ===== 10. Interview Q&A =====
h1("10. Domande probabili al colloquio — e risposte pronte")
qa = [
    ("Why did value of production fall more than revenue in 2024?",
     "Perché contiene la variazione dei lavori in corso: nel 2023 il WIP è cresciuto di +€19,3M (gonfiando il valore della produzione), nel 2024 si è smaltito (-€0,5M). È timing di commessa, non calo di domanda strutturale."),
    ("Over time or point in time for your machines?",
     "Dipende dalle clausole: per macchine su misura senza uso alternativo e con diritto al pagamento per il lavoro svolto, over time (IFRS 15.35c); altrimenti point in time alla consegna/collaudo. Da confermare contratto per contratto in DD."),
    ("How would you normalise EBITDA here?",
     "Depuro lo swing di WIP, one-off e componenti infragruppo (transfer pricing), normalizzo su più esercizi e su backlog/margine-a-finire; misuro la redditività stand-alone."),
    ("How do you control project margins?",
     "Cost-to-complete/EAC rivisti mensilmente, margin-at-completion per commessa, gestione varianti/claim, accantonamento immediato su commesse in perdita, disciplina di pricing sul backlog."),
    ("Are customer advances 'free cash'?",
     "No: sono contract liability, un'obbligazione di performance. Ottimi per il circolante (DSO basso) ma vanno gestiti su milestone; possibile componente finanziaria significativa."),
]
for q, a in qa:
    para([("Q: ", True, BLUE), (q, True)], after=1)
    para([("A: ", True, RED), (a, False)], after=4)

# ===== 11. Glossary =====
h1("11. Glossario rapido IT / EN")
table(["Italiano", "English", "Nota"],
      [["Lavori in corso su ordinazione", "Work in progress / contract asset", "OIC 23"],
       ["Percentuale di completamento", "Percentage-of-completion (over time)", "cost-to-cost"],
       ["Commessa completata", "Completed-contract", "se stime non attendibili"],
       ["Acconti da clienti", "Customer advances / contract liability", "passivo"],
       ["Costi a finire", "Cost-to-complete / EAC", "stima chiave"],
       ["Commessa in perdita", "Onerous contract", "accantonamento immediato"],
       ["Corrispettivo variabile", "Variable consideration", "penali/bonus, con constraint"],
       ["Obbligazione di fare", "Performance obligation (PO)", "unità di ricavo"]])

para([("Tip finale: ", True), ("entra in sala dicendo che “leggi” il loro 2024 in chiave di commessa — dimostri subito di aver capito il business e disinneschi il dubbio sul gap di project accounting.", False, RED)], before=4)

doc.save(OUT)
print("Saved:", OUT, "| paragraphs:", len(doc.paragraphs), "| tables:", len(doc.tables))
